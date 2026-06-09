from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/shop_afrik/Shop_Afrik_Project_Plan_and_Roadmap.docx"

COLORS = {
    "dark": "101418",
    "panel": "171D24",
    "teal": "0BA99D",
    "aqua": "3ED1C2",
    "deep_teal": "063F3B",
    "coral": "F56565",
    "muted": "9AA4B2",
    "border": "D9E2E7",
    "light_fill": "F3FBFA",
    "gray_fill": "F2F4F7",
    "ink": "1F2933",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_table_borders(table, color="D9E2E7"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_meta_row(table, key, value):
    row = table.add_row()
    set_cell_text(row.cells[0], key, bold=True, color=COLORS["deep_teal"])
    set_cell_text(row.cells[1], value, color=COLORS["ink"])
    row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(style="Body Text")
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, body, fill="F3FBFA"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.4)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_table_borders(table, color="BDEDE7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(COLORS["deep_teal"])
    r.font.size = Pt(11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for col, width in zip(table.columns, widths):
            col.width = Inches(width)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for i, h in enumerate(headers):
        set_cell_shading(header_row.cells[i], COLORS["deep_teal"])
        set_cell_text(header_row.cells[i], h, bold=True, color="FFFFFF", size=9.5)
        header_row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], str(value), color=COLORS["ink"], size=9.2)
            row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_table_borders(table)
    doc.add_paragraph()
    return table


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    body = styles["Body Text"]
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["deep_teal"], 16, 8),
        ("Heading 2", 13, COLORS["teal"], 12, 6),
        ("Heading 3", 12, COLORS["deep_teal"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.15


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Shop Afrik")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(COLORS["deep_teal"])

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Project Plan and Technical Roadmap")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(COLORS["teal"])

    add_body(
        doc,
        "A professional planning brief for building Shop Afrik as a full African e-commerce marketplace integrated with QR Wallet business payments.",
    )

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.5)
    add_meta_row(table, "Document type", "Product plan, architecture brief, and implementation roadmap")
    add_meta_row(table, "Prepared for", "Shop Afrik project kickoff")
    add_meta_row(table, "Version", "Planning draft 1.0")
    add_meta_row(table, "Date", "June 5, 2026")
    add_meta_row(table, "Source material", "Existing Shop Afrik specs, QR Wallet integration plans, and visual color reference")
    set_table_borders(table)
    doc.add_paragraph()

    add_callout(
        doc,
        "Planning Position",
        "Start with a focused marketplace MVP: buyer shopping, seller product management, admin approvals, QR Wallet pay-now checkout, order tracking, refunds, and day-8 seller settlement. Defer deep links and pay-on-delivery until the core marketplace flow is stable.",
    )
    doc.add_page_break()


def add_contents(doc):
    add_heading(doc, "1. Executive Summary", 1)
    add_body(
        doc,
        "Shop Afrik is planned as a full multi-seller e-commerce marketplace for African markets. The platform will support buyers, sellers, and administrators while using QR Wallet as the trusted payment rail. The first release should prove the core marketplace loop: discover a product, place an order, pay through QR Wallet, fulfill the order, handle refunds within a clear window, and settle sellers after the refund period.",
    )
    add_callout(
        doc,
        "Recommended MVP",
        "Build a mobile-first Flutter marketplace backed by a dedicated Shop Afrik Firebase project. Use QR Wallet merchant QR payment for v1, require seller KYC through QR Wallet, and hold proceeds in the Shop Afrik business wallet until day 8 after delivery.",
    )

    add_heading(doc, "2. Product Vision", 1)
    add_body(
        doc,
        "Shop Afrik should feel like a trusted African commerce platform: practical enough for everyday shopping, strong enough for seller operations, and disciplined enough for payments, refunds, and admin control. The product should eventually support multiple African countries, local currencies, QR Wallet payments, mobile money expansion, and country-specific logistics.",
    )
    add_bullets(
        doc,
        [
            "Buyers can discover products, compare options, pay securely, track orders, and request refunds.",
            "Sellers can onboard, upload products, manage stock, process orders, and monitor earnings.",
            "Admins can approve sellers and products, manage refunds, configure commissions, view reports, and investigate risk.",
            "QR Wallet remains the payment foundation while Shop Afrik owns the e-commerce experience.",
        ],
    )

    add_heading(doc, "3. Product Scope", 1)
    add_table(
        doc,
        ["Area", "MVP Scope", "Deferred"],
        [
            ("Buyer app", "Auth, catalog, product detail, cart, QR checkout, orders, refunds, wishlist, reviews", "Live chat, flash sales, comparison, price-drop alerts"),
            ("Seller dashboard", "Seller onboarding, KYC check, product CRUD, inventory, order updates, payout visibility", "Bulk CSV upload, advanced analytics, coupons, vacation mode"),
            ("Admin dashboard", "Seller/product approval, order overview, refund control, commission settings, audit logs", "A/B tests, advanced fraud scoring, warehouse operations"),
            ("Payments", "QR Wallet merchant QR pay-now flow, business wallet collection, refunds, day-8 settlement", "Deep link handoff, pay on delivery, non-wallet payment methods"),
            ("Logistics", "Basic delivery statuses and delivery fee model", "Courier app, live GPS dispatch, warehouse inspection tooling"),
        ],
        widths=[1.35, 3.1, 1.95],
    )

    add_heading(doc, "4. Platform Architecture", 1)
    add_body(
        doc,
        "The recommended architecture is a separate Shop Afrik app and Firebase backend integrated with QR Wallet through Cloud Functions. This keeps the wallet product stable while giving Shop Afrik freedom to evolve its marketplace data model, seller workflows, and admin controls.",
    )
    add_table(
        doc,
        ["Layer", "Responsibility"],
        [
            ("Shop Afrik Flutter app", "Buyer, seller, and admin experiences. Mobile-first, with admin web support when needed."),
            ("Shop Afrik Firebase project", "Products, categories, orders, buyers, sellers, reviews, refunds, notifications, audit logs, and settlement records."),
            ("Shop Afrik Cloud Functions", "Order creation, QR payload generation, payment confirmation, settlement jobs, refund workflow, seller/admin permissions."),
            ("QR Wallet Firebase project", "Business wallet, wallet balance, QR signing, sendMoney, refunds, holds, withdrawals, and existing KYC/wallet infrastructure."),
        ],
        widths=[2.1, 4.25],
    )
    add_callout(
        doc,
        "Architecture Decision",
        "Use the QR Wallet business wallet pattern, not a QR Wallet tenant refactor. The existing planning documents identify this as lower risk because it reuses established QR Wallet functions while keeping new Shop Afrik logic additive.",
    )

    add_heading(doc, "5. QR Wallet Payment Flow", 1)
    add_numbered(
        doc,
        [
            "Buyer adds items to cart in Shop Afrik and starts checkout.",
            "Shop Afrik calculates subtotal, payment fee, delivery fee, and total.",
            "Shop Afrik calls QR Wallet's signed QR payload function for the Shop Afrik business wallet.",
            "Buyer opens QR Wallet, scans the QR code, confirms payment, and enters PIN.",
            "QR Wallet sendMoney debits the buyer and credits the Shop Afrik business wallet.",
            "Shop Afrik detects the transaction, marks the order as paid, and begins fulfillment.",
            "After delivery and the 7-day refund window, Shop Afrik settles the seller share to the seller's QR Wallet.",
        ],
    )

    add_heading(doc, "6. Money, Refunds, and Settlement", 1)
    add_table(
        doc,
        ["Policy", "Recommendation"],
        [
            ("Commission", "Start at 15 percent, configurable globally and later per category."),
            ("Payment fee", "Buyer pays payment processing fee as a separate checkout line item."),
            ("Refund window", "7 days from courier-confirmed delivery."),
            ("Partial refunds", "Allowed for multi-item orders; single-item orders are full refund or no refund."),
            ("Settlement", "Day 8 auto-settlement to seller QR Wallet after refund window closes."),
            ("Seller KYC", "Required through QR Wallet before seller approval."),
            ("Admin escalation", "Use tiered refund approval based on refund amount and admin role."),
        ],
        widths=[1.7, 4.7],
    )

    add_heading(doc, "7. Core Data Model", 1)
    add_table(
        doc,
        ["Collection", "Purpose"],
        [
            ("products", "Catalog records with seller, price, images, stock, category, and approval status."),
            ("categories", "Category tree for browse, filters, and admin management."),
            ("orders", "Buyer order totals, payment status, delivery status, timestamps, and settlement status."),
            ("sellers", "Store profile, QR Wallet link, KYC status, rating, approval status, and payout settings."),
            ("buyers", "Customer profile, addresses, wishlist, and notification preferences."),
            ("reviews", "Verified-buyer product reviews and ratings."),
            ("refund_requests", "Buyer refund requests, evidence, inspection result, admin decisions, and escalation state."),
            ("settlements", "Day-8 payout records from Shop Afrik to sellers."),
            ("notifications", "In-app notification history for buyers, sellers, and admins."),
            ("admin_audit", "Append-only audit trail for admin and financial actions."),
        ],
        widths=[1.75, 4.65],
    )

    add_heading(doc, "8. Visual Direction", 1)
    add_body(
        doc,
        "The visual reference points toward a premium dark mobile interface with teal action states, soft aqua highlights, and coral danger states. Shop Afrik should borrow the trust and polish of fintech UI, then add stronger product imagery for the shopping experience.",
    )
    add_table(
        doc,
        ["Token", "Hex", "Usage"],
        [
            ("Primary dark", "#101418", "App background, high-trust surfaces"),
            ("Panel dark", "#171D24", "Cards, forms, product panels, bottom sheets"),
            ("Deep teal", "#063F3B", "Brand anchor, headers, dark accents"),
            ("Primary teal", "#0BA99D", "Main buttons, active tabs, payment confirmation"),
            ("Bright aqua", "#3ED1C2", "Highlights, success states, gradients"),
            ("Danger coral", "#F56565", "Refund warnings, delete actions, failed payment states"),
            ("Muted text", "#9AA4B2", "Secondary labels and helper text"),
        ],
        widths=[1.55, 1.15, 3.7],
    )
    add_bullets(
        doc,
        [
            "Use dark surfaces for account, payment, checkout, seller finance, and admin areas.",
            "Use product imagery generously on buyer screens so the app does not feel like a wallet with products added.",
            "Use teal for the main path forward and coral only for destructive, failed, or refund-related actions.",
            "Keep cards compact with moderate radius, strong spacing, and readable mobile typography.",
        ],
    )

    add_heading(doc, "9. Build Roadmap", 1)
    add_table(
        doc,
        ["Phase", "Outcome", "Key Work"],
        [
            ("Phase 0: Decisions", "Ready-to-build scope", "Choose countries, Firebase project ID, commission, refund thresholds, delivery model, and initial categories."),
            ("Phase 1: Project setup", "New app foundation", "Create GitHub repo, scaffold Flutter app, connect Firebase, define routing, theme, and base services."),
            ("Phase 2: Backend", "Marketplace data foundation", "Products, sellers, buyers, orders, payment confirmation, state machine, security rules, and Cloud Functions."),
            ("Phase 3: Buyer MVP", "Customer shopping flow", "Home, categories, search, product detail, cart, checkout, QR payment, orders, wishlist, reviews."),
            ("Phase 4: Seller MVP", "Seller operations", "Seller onboarding, product CRUD, inventory, order management, sales and payout visibility."),
            ("Phase 5: Admin MVP", "Platform control", "Approval queues, orders, refunds, commissions, audit logs, reports, and role controls."),
            ("Phase 6: Enhancements", "Smoother commerce", "Deep links, pay on delivery, courier confirmation, advanced notifications, promotions, analytics."),
        ],
        widths=[1.45, 1.7, 3.25],
    )

    add_heading(doc, "10. Open Decisions", 1)
    add_table(
        doc,
        ["Decision", "Recommended Default"],
        [
            ("Initial countries", "Start with Ghana and Nigeria if operations are ready; otherwise Ghana only for faster launch."),
            ("Firebase project", "Create a dedicated Shop Afrik Firebase project separate from QR Wallet."),
            ("Commission", "15 percent for MVP, stored as admin-configurable platform setting."),
            ("Refund Tier 1", "Admin can approve up to NGN 50,000 equivalent."),
            ("Refund Tier 2", "Admin plus supervisor approval up to NGN 300,000 equivalent."),
            ("Low stock threshold", "20 percent of initial stock or a seller-defined minimum quantity."),
            ("Delivery code", "Optional in MVP, recommended before scaling courier operations."),
            ("Minimum order amount", "Decide per country after delivery fee model is selected."),
            ("Initial categories", "Start with 10 to 15 main categories."),
            ("Max products per seller", "500 products per seller for v1."),
        ],
        widths=[2.1, 4.3],
    )

    add_heading(doc, "11. Immediate Next Steps", 1)
    add_numbered(
        doc,
        [
            "Approve this planning document or mark changes.",
            "Choose the final MVP country scope and Firebase project name.",
            "Create the GitHub repository for Shop Afrik.",
            "Create the Flutter project and commit the initial scaffold.",
            "Implement theme tokens, routing shell, Firebase setup, and app role structure.",
            "Build backend data model and security rules before UI screens rely on live data.",
        ],
    )

    add_heading(doc, "Appendix A: Source Planning Inputs", 1)
    add_bullets(
        doc,
        [
            "shop-afrik-complete-spec.curentmd.md",
            "shop-afrik-integration-plan-v3-FINAL.md",
            "shop-afrik-qr-wallet-integration-plan ONE.md.pdf",
            "Shop Afrik visual color reference image",
        ],
    )


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Shop Afrik Project Plan and Technical Roadmap")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def main():
    doc = Document()
    style_document(doc)
    add_title_page(doc)
    add_contents(doc)
    add_footer(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
